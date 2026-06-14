"""
Phase 6 — Enterprise template reconstruction engine.

Loads an uploaded .docx template (or default) and replaces all dynamic business
content while preserving fonts, styles, borders, spacing, and layout.

Template tables and narrative are placeholder only — values and paragraphs are
overwritten in-place from approved extraction and Phase 4 commentary.
"""

from __future__ import annotations

import copy
import logging
import re
from collections.abc import Callable
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Iterator

import pandas as pd
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from data.metric_aliases import (
    APPROVED_METRICS,
    NOT_DISCLOSED,
    TABLE1_PERIODS,
    TABLE2_PERIODS,
)
from services.normalizer import (
    canonicalize_table1_period,
    canonicalize_table2_period,
    format_crore_display,
    is_ratio_metric,
    normalize_text,
)
from services.reconstruction.similarity import score_row_match
from services.report_generator import _issuer_name_from_records
from services.review_manager import periods_from_records, split_records_by_table

logger = logging.getLogger("credit_review")

PLACEHOLDER_RE = re.compile(r"\{\{([A-Z0-9_]+)\}\}")

DEFAULT_TEMPLATE_FILENAME = "enterprise_default.docx"
FINAL_DOCX_NAME = "final_credit_review.docx"
FINAL_PDF_NAME = "final_credit_review.pdf"

# Mode A — heading keywords (normalized substring match)
SECTION_HEADING_MAP: dict[str, tuple[str, ...]] = {
    "yearly_financials": (
        "yearly financial",
        "annual financial",
        "yearly financials",
        "march year-end",
    ),
    "half_year_financials": (
        "half-year financial",
        "half year financial",
        "half-year financials",
        "h1 financial",
    ),
    "commentary": ("commentary",),
    "business_profile": ("business profile", "company profile", "issuer overview"),
    "profitability": ("profitability", "financial strength"),
    "capitalisation": ("capitalisation", "capitalization", "capital adequacy"),
    "liquidity": ("liquidity", "funding profile"),
    "company_profile": ("company profile", "issuer overview", "about the company"),
    "validation": ("validation", "validation notes"),
    "cio": ("cio", "fund manager", "investment view", "recommendation"),
}

# Process specific narrative sections before generic catch-alls.
SECTION_PROCESS_ORDER: tuple[str, ...] = (
    "yearly_financials",
    "half_year_financials",
    "business_profile",
    "profitability",
    "capitalisation",
    "liquidity",
    "company_profile",
    "validation",
    "cio",
    "commentary",
)

DATE_LINE_RE = re.compile(
    r"\b(?:as at|as on|review date|report date|dated)\b|\b\d{1,2}[\s/.-]"
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s/.-]\d{2,4}\b",
    re.IGNORECASE,
)


@dataclass
class ReportContext:
    """All injectable content for template replacement."""

    company_name: str
    report_date: str
    issuer_overview: str
    yearly_df: pd.DataFrame
    halfyear_df: pd.DataFrame
    commentary_yearly: list[str] = field(default_factory=list)
    commentary_halfyear: list[str] = field(default_factory=list)
    commentary_full: str = ""
    validation_notes: str = ""
    recommendation: str = (
        "Investment recommendation to be completed by the fund manager."
    )
    cio_content: str = (
        "Investment view and recommendation to be completed by the fund manager.\n\n"
        "Rating: _______________\n"
        "Outlook: _______________\n"
        "Key risks: _______________"
    )
    warnings: list[str] = field(default_factory=list)
    commentary_sections: list[dict[str, Any]] = field(default_factory=list)


def build_report_context(
    reviewed_records: list[dict[str, Any]],
    commentary: dict[str, Any],
    warnings: list[str],
) -> ReportContext:
    table1, table2 = split_records_by_table(reviewed_records)
    t1_periods = periods_from_records(table1) or TABLE1_PERIODS
    t2_periods = periods_from_records(table2) or TABLE2_PERIODS
    issuer = _issuer_name_from_records(reviewed_records)
    memo_sections = commentary.get("sections", [])
    section_paragraphs = [s.get("paragraph", "") for s in memo_sections]
    yearly_lines = section_paragraphs or commentary.get("yearly", {}).get("paragraphs", [])
    half_lines = section_paragraphs or commentary.get("half_year", {}).get("paragraphs", [])
    validation = (
        "\n".join(f"• {w}" for w in warnings)
        if warnings
        else "No validation warnings on approved values."
    )
    return ReportContext(
        company_name=issuer,
        report_date=datetime.now().strftime("%d %B %Y"),
        issuer_overview=(
            f"This credit review for {issuer} uses disclosed standalone annual report "
            "and investor presentation data. Monetary figures are ₹ crore unless "
            "noted; ratios are in percent. 'Not disclosed' means not explicitly found."
        ),
        yearly_df=pivot_template_table(table1, t1_periods),
        halfyear_df=pivot_template_table(table2, t2_periods),
        commentary_yearly=list(yearly_lines),
        commentary_halfyear=list(half_lines),
        commentary_full=commentary.get("full_text", ""),
        commentary_sections=memo_sections,
        validation_notes=validation,
        warnings=list(warnings),
    )


def _metric_to_slug(metric: str) -> str:
    return re.sub(r"[^A-Z0-9]+", "_", metric.upper()).strip("_")


def _slugs_for_metric(metric: str) -> tuple[str, ...]:
    """Placeholder suffixes used in enterprise templates (e.g. Y1_TOTAL_INCOME)."""
    primary = _metric_to_slug(metric)
    aliases: dict[str, tuple[str, ...]] = {
        "CAPITAL_ADEQUACY_RATIO": ("CAR", "CRAR"),
        "TIER_I_CAPITAL_RATIO": ("TIER_I", "TIER_1", "TIER_1_CAPITAL_RATIO"),
    }
    extra = aliases.get(primary, ())
    return (primary, *extra)


def build_full_placeholder_map(
    ctx: ReportContext,
    yearly_lookup: dict[tuple[str, str], str],
    halfyear_lookup: dict[tuple[str, str], str],
    t1_periods: tuple[str, ...],
    t2_periods: tuple[str, ...],
) -> dict[str, str]:
    """
    All {{TOKEN}} values for enterprise templates.

    Supports per-cell tokens such as {{ISSUER}}, {{Y1_DATE}}, {{Y1_PAT}},
    {{H1_PERIOD}}, {{H1_NII}}, etc.
    """
    values = _placeholder_values(ctx)
    values["ISSUER"] = ctx.company_name
    values["NIC_CODE"] = "—"
    values["SECTOR"] = "—"
    values["PREPARED_BY"] = "—"
    values["REVIEWED_BY"] = "—"
    if t1_periods and t2_periods:
        values["REVIEW_PERIOD"] = f"Year ended {t1_periods[0]}; Half-year {t2_periods[0]}"
    elif t1_periods:
        values["REVIEW_PERIOD"] = str(t1_periods[0])
    else:
        values["REVIEW_PERIOD"] = ctx.report_date

    for i, period in enumerate(t1_periods[:3], start=1):
        values[f"Y{i}_DATE"] = period
        for metric in APPROVED_METRICS:
            val = yearly_lookup.get((metric, period), NOT_DISCLOSED)
            for slug in _slugs_for_metric(metric):
                values[f"Y{i}_{slug}"] = val

    for i, period in enumerate(t2_periods[:2], start=1):
        values[f"H{i}_PERIOD"] = period
        for metric in APPROVED_METRICS:
            val = halfyear_lookup.get((metric, period), NOT_DISCLOSED)
            for slug in _slugs_for_metric(metric):
                values[f"H{i}_{slug}"] = val

    return values


def build_approved_value_lookup(
    records: list[dict[str, Any]],
) -> dict[tuple[str, str], str]:
    """Map (metric, period) → clean display string for template cells (no page hints)."""
    lookup: dict[tuple[str, str], str] = {}
    for rec in records:
        metric = rec.get("metric", "")
        period = rec.get("period", "")
        val = rec.get("approved_value")
        if val is None:
            lookup[(metric, period)] = NOT_DISCLOSED
            continue
        try:
            num = float(val)
        except (TypeError, ValueError):
            lookup[(metric, period)] = NOT_DISCLOSED
            continue
        text = format_crore_display(num)
        if is_ratio_metric(metric):
            text = f"{text}%"
        lookup[(metric, period)] = text
    return lookup


def pivot_template_table(
    records: list[dict[str, Any]],
    periods: tuple[str, ...],
) -> pd.DataFrame:
    """Pivot for template injection — clean values only, no page references."""
    lookup = build_approved_value_lookup(records)
    rows: list[dict[str, Any]] = []
    for metric in APPROVED_METRICS:
        row: dict[str, Any] = {"Metric": metric}
        for period in periods:
            row[period] = lookup.get((metric, period), NOT_DISCLOSED)
        rows.append(row)
    return pd.DataFrame(rows)


def _split_heading_inline(text: str) -> tuple[str, str | None]:
    """Split 'Capitalisation: body text' into heading label and inline body."""
    raw = (text or "").strip()
    if ":" not in raw:
        return raw, None
    head, _, body = raw.partition(":")
    if len(head.strip()) > 80:
        return raw, None
    return head.strip() + ":", body.strip() if body.strip() else ""


def iter_all_tables(doc: Document) -> Iterator[Table]:
    """Yield every table in body, headers, footers, and nested table cells."""

    def _walk_tables(tables: list[Table]) -> Iterator[Table]:
        for table in tables:
            yield table
            for row in table.rows:
                for cell in row.cells:
                    if cell.tables:
                        yield from _walk_tables(cell.tables)

    yield from _walk_tables(list(doc.tables))
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            yield from _walk_tables(list(part.tables))


def _table_header_row_index(table: Table) -> tuple[int, dict[int, str], str] | None:
    """
    Detect financial table type from header row(s).

    Returns (header_row_index, {col_index: canonical_period}, table_type).
    """
    best: tuple[int, dict[int, str], str, int] | None = None

    for row_idx in range(min(3, len(table.rows))):
        yearly_map: dict[int, str] = {}
        half_map: dict[int, str] = {}
        for col_idx, cell in enumerate(table.rows[row_idx].cells):
            if col_idx == 0:
                continue
            header_text = cell.text.strip()
            if not header_text:
                continue
            p1 = canonicalize_table1_period(header_text)
            if p1:
                yearly_map[col_idx] = p1
            p2 = canonicalize_table2_period(header_text)
            if p2:
                half_map[col_idx] = p2

        if len(yearly_map) >= len(half_map) and len(yearly_map) >= 1:
            score = len(yearly_map)
            if best is None or score > best[3]:
                best = (row_idx, yearly_map, "yearly", score)
        elif len(half_map) >= 1:
            score = len(half_map)
            if best is None or score > best[3]:
                best = (row_idx, half_map, "half_year", score)

    if best is None:
        return None
    row_idx, col_map, table_type, _ = best
    return row_idx, col_map, table_type


def _match_row_metric(label: str) -> str | None:
    from data.metric_aliases import METRIC_ALIASES
    if not label or not label.strip():
        return None

    # Direct hardcoded mappings for common template variants
    DIRECT_MAP = {
        "crar": "Capital Adequacy Ratio",
        "crar (%)": "Capital Adequacy Ratio",
        "crar – tier i (%)": "Tier I Capital Ratio",
        "crar - tier i (%)": "Tier I Capital Ratio",
        "car (%)": "Capital Adequacy Ratio",
        "car": "Capital Adequacy Ratio",
        "gnpa (%)": "GNPA",
        "nnpa (%)": "NNPA",
        "roa (%)": "ROA",
        "roe (%)": "ROE",
        "customer financial assets": "Advances",
        "net worth": "Net Worth",
    }

    from services.normalizer import normalize_text
    norm = normalize_text(label)
    if norm in DIRECT_MAP:
        return DIRECT_MAP[norm]

    # First pass: check full alias list for exact or near-exact match
    from services.normalizer import normalize_text
    norm_label = normalize_text(label)
    for metric, aliases in METRIC_ALIASES.items():
        for alias in aliases:
            if normalize_text(alias) == norm_label:
                return metric
            # Handle em-dash vs hyphen and spacing variants
            if normalize_text(alias).replace("–", "-").replace("—", "-") == \
               norm_label.replace("–", "-").replace("—", "-"):
                return metric

    # Second pass: fuzzy score fallback
    best_metric: str | None = None
    best_score = 0.0
    for metric in APPROVED_METRICS:
        score, _, _ = score_row_match(metric, label)
        if score > best_score:
            best_score = score
            best_metric = metric
    return best_metric if best_score >= 0.75 else None


def overwrite_financial_table_in_place(
    table: Table,
    lookup: dict[tuple[str, str], str],
    *,
    table_type: str | None = None,
) -> bool:
    """
    Replace data cells in an existing financial table; preserve cell formatting.

    Matches rows by metric label and columns by period header. Clears stale values
    in period columns when the row maps to an approved metric.
    """
    detected = _table_header_row_index(table)
    if detected is None:
        return False
    header_row_idx, col_period_map, detected_type = detected
    if table_type and detected_type != table_type:
        return False

    changed = False
    for row_idx in range(header_row_idx + 1, len(table.rows)):
        row = table.rows[row_idx]
        if not row.cells:
            continue
        label = row.cells[0].text.strip()
        metric = _match_row_metric(label)
        if metric is None:
            continue
        for col_idx, period in col_period_map.items():
            if col_idx >= len(row.cells):
                continue
            new_val = lookup.get((metric, period), NOT_DISCLOSED)
            old_val = row.cells[col_idx].text.strip()
            if old_val != new_val:
                _set_cell_text_preserve(row.cells[col_idx], new_val)
                changed = True
    return changed


def replace_all_financial_tables(
    doc: Document,
    yearly_lookup: dict[tuple[str, str], str],
    halfyear_lookup: dict[tuple[str, str], str],
) -> list[str]:
    """Scan every table and overwrite financial data in-place."""
    log: list[str] = []
    seen_ids: set[int] = set()
    for table in iter_all_tables(doc):
        tbl_id = id(table._tbl)
        if tbl_id in seen_ids:
            continue
        seen_ids.add(tbl_id)

        detected = _table_header_row_index(table)
        if detected is None:
            continue
        _, _, table_type = detected
        lookup = yearly_lookup if table_type == "yearly" else halfyear_lookup
        if overwrite_financial_table_in_place(table, lookup, table_type=table_type):
            log.append(f"Replaced {table_type} financial table values in-place")
    return log


def replace_title_metadata(doc: Document, ctx: ReportContext) -> list[str]:
    """Replace company name and review date in title block (preserve layout)."""
    log: list[str] = []
    company_set = False

    for i, paragraph in enumerate(doc.paragraphs[:15]):
        text = paragraph.text.strip()
        if not text:
            continue
        if "{{" in text:
            if _replace_in_paragraph(paragraph, {"ISSUER": ctx.company_name, "ISSUER_NAME": ctx.company_name, "COMPANY_NAME": ctx.company_name}):
                log.append("Replaced issuer placeholder in title block")
                company_set = True
            continue
        if _is_heading_paragraph(paragraph) and i > 0:
            break
        if DATE_LINE_RE.search(text):
            _set_paragraph_text_preserve_style(paragraph, ctx.report_date)
            log.append("Replaced review date in title block")
            continue
        if (
            not company_set
            and i <= 4
            and len(text) < 120
            and normalize_text(text) != normalize_text(ctx.company_name)
            and not re.search(r"credit review|annual review|financial review", text, re.I)
        ):
            _set_paragraph_text_preserve_style(paragraph, ctx.company_name)
            company_set = True
            log.append("Replaced company name in title block")

    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text.strip()
        if not text:
            continue
        if "{{" in text:
            continue
        if DATE_LINE_RE.search(text) and normalize_text(text) != normalize_text(ctx.report_date):
            _set_paragraph_text_preserve_style(paragraph, ctx.report_date)
            log.append("Replaced date field")
    return log


def _placeholder_values(ctx: ReportContext) -> dict[str, str]:
    return {
        "COMPANY_NAME": ctx.company_name,
        "ISSUER_NAME": ctx.company_name,
        "REPORT_DATE": ctx.report_date,
        "DATE": ctx.report_date,
        "ISSUER_OVERVIEW": ctx.issuer_overview,
        "COMMENTARY": ctx.commentary_full,
        "COMMENTARY_YEARLY": "\n".join(ctx.commentary_yearly),
        "COMMENTARY_HALFYEAR": "\n".join(ctx.commentary_halfyear),
        "COMMENTARY_HALF_YEAR": "\n".join(ctx.commentary_halfyear),
        "VALIDATION_NOTES": ctx.validation_notes,
        "RECOMMENDATION": ctx.recommendation,
        "CIO_FUND_MANAGER": ctx.cio_content,
        "CIO_CONTENT": ctx.cio_content,
    }


def _table_placeholders() -> frozenset[str]:
    return frozenset({"YEARLY_TABLE", "HALFYEAR_TABLE", "HALF_YEAR_TABLE"})


def iter_all_paragraphs(doc: Document) -> Iterator[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            for p in part.paragraphs:
                yield p
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


def _normalize_heading(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip().lower())


def _is_heading_paragraph(paragraph: Paragraph) -> bool:
    style = (paragraph.style.name or "").lower() if paragraph.style else ""
    if "heading" in style:
        return True
    text = paragraph.text.strip()
    if not text or len(text) > 120:
        return False
    if paragraph.runs and paragraph.runs[0].bold and len(text) < 80:
        return True
    return False


def _set_paragraph_text_preserve_style(paragraph: Paragraph, text: str) -> None:
    """Replace visible text; keep paragraph style and first run character formatting."""
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_in_paragraph(paragraph: Paragraph, replacements: dict[str, str]) -> bool:
    full = paragraph.text
    if not full or "{{" not in full:
        return False
    new_text = full
    changed = False
    for key, val in replacements.items():
        token = f"{{{{{key}}}}}"
        if token in new_text:
            new_text = new_text.replace(token, val)
            changed = True
    if changed:
        _set_paragraph_text_preserve_style(paragraph, new_text)
    return changed


def _set_cell_text_preserve(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.text = str(text)
        return
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = str(text)
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = str(text)


def _clone_table_row(table: Table, template_row_index: int = -1) -> None:
    template_tr = table.rows[template_row_index]._tr
    new_tr = copy.deepcopy(template_tr)
    table._tbl.append(new_tr)


def _fill_table_from_dataframe(table: Table, df: pd.DataFrame) -> None:
    """Populate table cells; expand rows using last row style when needed."""
    if df.empty:
        return
    n_cols = len(df.columns)
    n_data = len(df)
    needed_rows = 1 + n_data

    while len(table.rows) < needed_rows:
        src = 1 if len(table.rows) > 1 else 0
        _clone_table_row(table, template_row_index=src)

    while len(table.rows) > needed_rows and len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    for j, col in enumerate(df.columns):
        if j < len(table.rows[0].cells):
            _set_cell_text_preserve(table.rows[0].cells[j], str(col))

    for i in range(n_data):
        row = df.iloc[i]
        for j, col in enumerate(df.columns):
            if j < len(table.rows[i + 1].cells):
                _set_cell_text_preserve(table.rows[i + 1].cells[j], str(row[col]))


def _insert_table_after_paragraph(
    doc: Document,
    paragraph: Paragraph,
    df: pd.DataFrame,
    *,
    style_name: str | None = None,
) -> Table:
    """Insert a new table after a paragraph; try to match an existing table style."""
    n_rows = max(2, len(df) + 1)
    n_cols = max(1, len(df.columns))
    table = doc.add_table(rows=n_rows, cols=n_cols)
    if style_name:
        try:
            table.style = style_name
        except Exception:
            pass
    elif doc.tables:
        try:
            table.style = doc.tables[0].style
        except Exception:
            pass
    paragraph._p.addnext(table._tbl)
    _fill_table_from_dataframe(table, df)
    return table


def _paragraph_in_table_cell(paragraph: Paragraph) -> bool:
    el = paragraph._element
    while el is not None:
        if el.tag.endswith("tbl"):
            return True
        el = el.getparent()
    return False


def _paragraph_table_placeholder(paragraph: Paragraph) -> str | None:
    text = paragraph.text.strip().upper()
    for name in _table_placeholders():
        token = f"{{{{{name}}}}}"
        if token in text or text == name:
            return name
    return None


def apply_placeholders(
    doc: Document,
    values: dict[str, str],
    *,
    yearly_df: pd.DataFrame | None = None,
    halfyear_df: pd.DataFrame | None = None,
) -> list[str]:
    """Mode B — replace {{PLACEHOLDER}} tokens in paragraphs and table cells."""
    log: list[str] = []
    table_style = doc.tables[0].style if doc.tables else None

    for paragraph in list(iter_all_paragraphs(doc)):
        tbl_key = _paragraph_table_placeholder(paragraph)
        if tbl_key:
            df = yearly_df if tbl_key == "YEARLY_TABLE" else halfyear_df
            if df is None:
                continue
            if _paragraph_in_table_cell(paragraph):
                lines = ["\t".join(str(c) for c in df.columns)]
                for _, row in df.iterrows():
                    lines.append("\t".join(str(row[c]) for c in df.columns))
                _set_paragraph_text_preserve_style(paragraph, "\n".join(lines))
                log.append(f"Filled cell text for {{{{{tbl_key}}}}}")
            else:
                _set_paragraph_text_preserve_style(paragraph, "")
                _insert_table_after_paragraph(
                    doc,
                    paragraph,
                    df,
                    style_name=table_style,
                )
                log.append(f"Inserted table for {{{{{tbl_key}}}}}")
            continue

        if _replace_in_paragraph(paragraph, values):
            preview = paragraph.text[:60].replace("\n", " ")
            log.append(f"Replaced placeholders in: {preview}…")

    return log


def _replace_placeholders_in_context(
    doc: Document,
    ctx: ReportContext,
    reviewed_records: list[dict[str, Any]],
) -> list[str]:
    """Replace all {{TOKEN}} placeholders including Y1_/H1_ financial cells."""
    table1, table2 = split_records_by_table(reviewed_records)
    t1_periods = periods_from_records(table1) or TABLE1_PERIODS
    t2_periods = periods_from_records(table2) or TABLE2_PERIODS
    yearly_lookup = build_approved_value_lookup(table1)
    halfyear_lookup = build_approved_value_lookup(table2)
    values = build_full_placeholder_map(
        ctx, yearly_lookup, halfyear_lookup, t1_periods, t2_periods,
    )
    return apply_placeholders(
        doc, values, yearly_df=ctx.yearly_df, halfyear_df=ctx.halfyear_df,
    )


def _body_paragraphs(doc: Document) -> list[Paragraph]:
    return list(doc.paragraphs)


def _find_section_range(
    paragraphs: list[Paragraph],
    keywords: tuple[str, ...],
) -> tuple[int, int] | None:
    start = None
    for i, p in enumerate(paragraphs):
        norm = _normalize_heading(p.text)
        if not norm:
            continue
        if any(kw in norm for kw in keywords):
            start = i
            break
    if start is None:
        return None
    end = len(paragraphs)
    for j in range(start + 1, len(paragraphs)):
        if _is_heading_paragraph(paragraphs[j]) and paragraphs[j].text.strip():
            end = j
            break
    return start, end


def _replace_section_paragraphs(
    paragraphs: list[Paragraph],
    start: int,
    end: int,
    lines: list[str],
) -> None:
    """
    Replace narrative body under a section heading.

    Clears all stale paragraphs in the range, then writes new content into the
    first available body slot(s). Handles inline 'Heading: body' on one line.
    """
    heading = paragraphs[start]
    title, inline_body = _split_heading_inline(heading.text)
    body_lines = [ln.strip() for ln in lines if ln and ln.strip()]

    if inline_body is not None:
        _set_paragraph_text_preserve_style(heading, title)

    for idx in range(start + 1, end):
        _set_paragraph_text_preserve_style(paragraphs[idx], "")

    if not body_lines:
        return

    if inline_body is not None or end - start <= 1:
        combined = " ".join(body_lines)
        if inline_body is not None and title:
            _set_paragraph_text_preserve_style(heading, f"{title} {combined}".strip())
        elif start + 1 < end:
            _set_paragraph_text_preserve_style(paragraphs[start + 1], combined)
        elif title:
            _set_paragraph_text_preserve_style(heading, f"{title} {combined}".strip())
        return

    for offset, text in enumerate(body_lines):
        idx = start + 1 + offset
        if idx < end:
            _set_paragraph_text_preserve_style(paragraphs[idx], text)
        else:
            break


def _table_after_paragraph_index(doc: Document, para_index: int) -> Table | None:
    """Return first table element immediately following paragraph at index."""
    if para_index < 0 or para_index >= len(doc.paragraphs):
        return None
    p_el = doc.paragraphs[para_index]._p
    nxt = p_el.getnext()
    if nxt is not None and nxt.tag.endswith("tbl"):
        for table in doc.tables:
            if table._tbl is nxt:
                return table
    for table in doc.tables:
        if table._tbl.getprevious() is p_el:
            return table
    return None


def apply_smart_sections(
    doc: Document,
    ctx: ReportContext,
    *,
    yearly_lookup: dict[tuple[str, str], str],
    halfyear_lookup: dict[tuple[str, str], str],
) -> list[str]:
    """Replace dynamic narrative and financial content under recognized headings."""
    paragraphs = _body_paragraphs(doc)
    log: list[str] = []
    handled_ranges: set[tuple[int, int]] = set()

    section_by_id = {
        s.get("id", ""): s.get("paragraph", "")
        for s in getattr(ctx, "commentary_sections", []) or []
    }

    commentary_blocks = [
        section_by_id[sid]
        for sid, _title in (
            ("business_profile", "Business Profile"),
            ("profitability", "Profitability"),
            ("capitalisation", "Capitalisation"),
            ("liquidity", "Liquidity"),
        )
        if section_by_id.get(sid)
    ]

    section_content: dict[str, Any] = {
        "yearly_financials": ctx.yearly_df,
        "half_year_financials": ctx.halfyear_df,
        "commentary": commentary_blocks or ([ctx.commentary_full] if ctx.commentary_full else []),
        "business_profile": [section_by_id.get("business_profile", ctx.issuer_overview)],
        "profitability": [section_by_id.get("profitability", "")],
        "capitalisation": [section_by_id.get("capitalisation", "")],
        "liquidity": [section_by_id.get("liquidity", "")],
        "company_profile": [section_by_id.get("business_profile", ctx.issuer_overview)],
        "validation": [ctx.validation_notes],
        "cio": [ctx.cio_content, ctx.recommendation],
    }

    for section_key in SECTION_PROCESS_ORDER:
        keywords = SECTION_HEADING_MAP.get(section_key)
        if not keywords:
            continue
        rng = _find_section_range(paragraphs, keywords)
        if rng is None:
            continue
        start, end = rng
        if (start, end) in handled_ranges:
            continue
        handled_ranges.add((start, end))
        log.append(f"Section matched: {section_key} @ paragraph {start}")

        if section_key in ("yearly_financials", "half_year_financials"):
            table = _table_after_paragraph_index(doc, start)
            lookup = (
                yearly_lookup
                if section_key == "yearly_financials"
                else halfyear_lookup
            )
            if table is not None:
                if overwrite_financial_table_in_place(
                    table,
                    lookup,
                    table_type=(
                        "yearly"
                        if section_key == "yearly_financials"
                        else "half_year"
                    ),
                ):
                    log.append(f"Replaced table under {section_key}")
                else:
                    df = section_content[section_key]
                    _fill_table_from_dataframe(table, df)
                    log.append(f"Filled table structure under {section_key}")
            else:
                _replace_section_paragraphs(
                    paragraphs,
                    start,
                    end,
                    [f"Financial data for {section_key.replace('_', ' ')}."],
                )
        elif section_key == "commentary":
            _replace_section_paragraphs(
                paragraphs, start, end, section_content["commentary"]
            )
            log.append("Replaced commentary section body")
        elif section_key in (
            "business_profile",
            "profitability",
            "capitalisation",
            "liquidity",
            "company_profile",
        ):
            text = section_by_id.get(section_key) or (
                section_content[section_key][0] if section_content[section_key] else ""
            )
            if text:
                _replace_section_paragraphs(paragraphs, start, end, [text])
                log.append(f"Replaced {section_key} narrative")
        elif section_key == "validation":
            _replace_section_paragraphs(
                paragraphs, start, end, [ctx.validation_notes]
            )
            log.append("Replaced validation notes")
        elif section_key == "cio":
            _replace_section_paragraphs(
                paragraphs, start, end, [ctx.cio_content, ctx.recommendation]
            )
            log.append("Replaced CIO / recommendation block")

    return log


def create_default_enterprise_template(path: Path) -> Path:
    """Create internal default template with explicit placeholders."""
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("Credit Review Report")
    doc.add_paragraph("{{COMPANY_NAME}}")
    doc.add_paragraph("{{DATE}}")
    doc.add_page_break()
    doc.add_heading("Company Profile", level=1)
    doc.add_paragraph("{{ISSUER_OVERVIEW}}")
    doc.add_heading("Yearly Financials", level=1)
    doc.add_paragraph("{{YEARLY_TABLE}}")
    doc.add_heading("Half-Year Financials", level=1)
    doc.add_paragraph("{{HALFYEAR_TABLE}}")
    doc.add_heading("Commentary", level=1)
    doc.add_paragraph("{{COMMENTARY}}")
    doc.add_heading("Validation Notes", level=1)
    doc.add_paragraph("{{VALIDATION_NOTES}}")
    doc.add_heading("CIO / Fund Manager", level=1)
    doc.add_paragraph("{{CIO_FUND_MANAGER}}")
    doc.add_paragraph("{{RECOMMENDATION}}")
    doc.save(str(path))
    logger.info("Default enterprise template created at %s", path)
    return path


def export_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """Best-effort PDF export (Windows: docx2pdf or Word COM)."""
    docx_path = Path(docx_path)
    pdf_path = Path(pdf_path)
    try:
        import docx2pdf  # type: ignore[import-untyped]

        docx2pdf.convert(str(docx_path), str(pdf_path))
        if pdf_path.is_file():
            logger.info("PDF exported via docx2pdf: %s", pdf_path)
            return True
    except Exception as exc:
        logger.debug("docx2pdf failed: %s", exc)

    try:
        import win32com.client  # type: ignore[import-untyped]

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
        doc.Close()
        word.Quit()
        if pdf_path.is_file():
            logger.info("PDF exported via Word COM: %s", pdf_path)
            return True
    except Exception as exc:
        logger.debug("Word COM PDF export failed: %s", exc)

    logger.warning("PDF export unavailable — DOCX output is still valid.")
    return False


def apply_template_reconstruction(
    doc: Document,
    ctx: ReportContext,
    reviewed_records: list[dict[str, Any]],
    *,
    on_status: Callable[[str], None] | None = None,
) -> dict[str, list[str]]:
    """
    Replace all dynamic business content while preserving template design.

    Order: placeholders → title metadata → financial tables (in-place) →
    narrative sections (clear stale text, inject new commentary).
    """

    def _status(msg: str) -> None:
        if on_status:
            on_status(msg)

    table1, table2 = split_records_by_table(reviewed_records)
    yearly_lookup = build_approved_value_lookup(table1)
    halfyear_lookup = build_approved_value_lookup(table2)

    _status("Applying enterprise report format…")
    placeholders = _replace_placeholders_in_context(doc, ctx, reviewed_records)
    _status("Updating issuer name and review period…")
    metadata = replace_title_metadata(doc, ctx)
    _status("Replacing financial tables…")
    tables = replace_all_financial_tables(doc, yearly_lookup, halfyear_lookup)
    _status("Injecting approved commentary…")
    sections = apply_smart_sections(
        doc,
        ctx,
        yearly_lookup=yearly_lookup,
        halfyear_lookup=halfyear_lookup,
    )
    _status("Formatting final report…")
    return {
        "placeholders": placeholders,
        "metadata": metadata,
        "tables": tables,
        "sections": sections,
    }


def format_enterprise_report(
    *,
    template_path: Path,
    reviewed_records: list[dict[str, Any]],
    commentary: dict[str, Any],
    warnings: list[str],
    output_dir: Path,
    on_status: Callable[[str], None] | None = None,
    llm_sections: dict[str, str] | None = None,
) -> dict[str, Any]:
    """
    Load template, inject content, save final_credit_review.docx (+ PDF if possible).
    """

    def _status(msg: str) -> None:
        if on_status:
            on_status(msg)

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_out = output_dir / FINAL_DOCX_NAME
    pdf_out = output_dir / FINAL_PDF_NAME

    _status("Loading enterprise template…")
    ctx = build_report_context(reviewed_records, commentary, warnings)
    doc = Document(str(template_path))

    recon_log = apply_template_reconstruction(
        doc, ctx, reviewed_records, on_status=on_status
    )

    for para in doc.paragraphs:
        text = para.text.strip()

        # Replace Company Profile placeholder
        if "company profile" in text.lower() or (
            "assessed based on disclosed" in text.lower()
        ):
            if llm_sections and llm_sections.get("company_profile"):
                para.clear()
                bold_run = para.add_run("Company Profile: ")
                bold_run.bold = True
                para.add_run(llm_sections["company_profile"])
            continue

        # Replace Comments date placeholders
        import re
        if re.match(r"^\d{1,2}\s+\w+\s+\d{4}$", text):
            para.clear()
            continue

        # Replace Recommendation placeholder
        if ("recommendation" in text.lower() and
            "fund manager" in text.lower()):
            if llm_sections and llm_sections.get("recommendation"):
                para.clear()
                bold_run = para.add_run("Recommendation: ")
                bold_run.bold = True
                para.add_run(llm_sections["recommendation"])
            continue

    # Write Profitability, Asset Quality, Capitalisation, Liquidity
    # after the Comments heading
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == "comments:":
            section_order = [
                ("profitability",  "Profitability"),
                ("asset_quality",  "Asset Quality"),
                ("capitalisation", "Capitalisation"),
                ("liquidity",      "Liquidity"),
            ]
            insert_after = para._element
            for key, label in reversed(section_order):
                text = llm_sections.get(key, "") if llm_sections else ""
                if not text:
                    continue
                from docx.oxml import OxmlElement
                new_para = OxmlElement("w:p")
                insert_after.addnext(new_para)
                from docx.text.paragraph import Paragraph
                new_p = Paragraph(new_para, para._element.getparent())
                bold_run = new_p.add_run(f"{label}: ")
                bold_run.bold = True
                new_p.add_run(text)
            break

    _status("Preparing DOCX export…")
    doc.save(str(docx_out))
    _status("Exporting PDF (if available)…")
    pdf_ok = export_docx_to_pdf(docx_out, pdf_out)

    flat_log = (
        recon_log.get("placeholders", [])
        + recon_log.get("metadata", [])
        + recon_log.get("tables", [])
        + recon_log.get("sections", [])
    )

    return {
        "docx_path": str(docx_out),
        "pdf_path": str(pdf_out) if pdf_ok else None,
        "template_used": str(template_path),
        "placeholder_log": flat_log,
        "section_log": recon_log.get("sections", []),
        "reconstruction_log": recon_log,
        "pdf_exported": pdf_ok,
    }


def resolve_template_path(
    templates_dir: Path,
    uploaded_bytes: bytes | None,
    uploaded_name: str | None,
) -> Path:
    """Use uploaded template or create/load default."""
    if uploaded_bytes:
        safe = re.sub(r"[^\w.\- ]", "_", uploaded_name or "uploaded_template.docx")
        if not safe.lower().endswith(".docx"):
            safe += ".docx"
        path = templates_dir / safe
        path.write_bytes(uploaded_bytes)
        return path

    default = templates_dir / DEFAULT_TEMPLATE_FILENAME
    if not default.is_file():
        create_default_enterprise_template(default)
    return default
