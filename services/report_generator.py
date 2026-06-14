"""
Phase 5 — Credit review DOCX report (AU Small Finance Bank style layout).

Uses python-docx. Values come only from approved extraction + Phase 4 commentary.
"""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from collections.abc import Callable
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement

from data.metric_aliases import NOT_DISCLOSED, TABLE1_PERIODS, TABLE2_PERIODS
from services.normalizer import format_crore_display, is_ratio_metric
from services.review_manager import (
    periods_from_records,
    pivot_review_table,
    split_records_by_table,
)

logger = logging.getLogger("credit_review")

FONT_NAME = "Calibri"
FONT_SIZE_PT = 11
HEADING_COLOR = RGBColor(0x00, 0x33, 0x66)


def _set_run_font(run, *, bold: bool = False, size_pt: int = FONT_SIZE_PT) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, bold=True, size_pt=FONT_SIZE_PT + (2 if level == 1 else 1))
    run.font.color.rgb = HEADING_COLOR


def _add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    _set_run_font(run)


def _add_table_from_dataframe(doc: Document, df, title: str) -> None:
    _add_heading(doc, title, level=2)
    if df.empty:
        _add_body(doc, "No data available.")
        return

    cols = list(df.columns)
    table = doc.add_table(rows=1 + len(df), cols=len(cols))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr_cells[i].text = str(col)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                _set_run_font(r, bold=True)

    for row_idx, (_, row) in enumerate(df.iterrows(), start=1):
        for col_idx, col in enumerate(cols):
            table.rows[row_idx].cells[col_idx].text = str(row[col])


def _issuer_name_from_records(records: list[dict[str, Any]]) -> str:
    # Check for explicit issuer override in session or records
    for rec in records:
        source = rec.get("source_filename") or rec.get("source_file") or ""
        if "kotak mahindra bank" in source.lower():
            return "Kotak Mahindra Bank Limited"
        if "kotak-mahindra-bank" in source.lower():
            return "Kotak Mahindra Bank Limited"

    for rec in records:
        explicit = rec.get("issuer_name") or rec.get("company_name")
        if explicit and str(explicit).strip():
            return str(explicit).strip()
    for rec in records:
        fn = rec.get("source_filename") or rec.get("source_file") or ""
        if fn:
            name = Path(fn).stem
            # Strip trailing year tokens like "BANK 25"
            parts = name.replace("_", " ").split()
            cleaned: list[str] = []
            for p in parts:
                if p.isdigit() or (len(p) == 2 and p.isdigit()):
                    continue
                if p.lower() in {"fy25", "fy24", "fy26", "ppt", "pdf"}:
                    continue
                cleaned.append(p)
            if cleaned:
                return " ".join(cleaned).title()
    return "Issuer"


def _provenance_rows(records: list[dict[str, Any]]) -> list[list[str]]:
    rows: list[list[str]] = []
    for rec in records:
        val = rec.get("approved_value")
        if val is None:
            display = NOT_DISCLOSED
        elif is_ratio_metric(rec["metric"]):
            display = f"{format_crore_display(float(val))}%"
        else:
            display = f"₹{format_crore_display(float(val))} crore"
        rows.append(
            [
                rec["metric"],
                rec["period"],
                display,
                str(rec.get("original_unit", "")),
                str(rec.get("page_number") or "—"),
                str(rec.get("source_document", "")),
                str(rec.get("source_filename") or rec.get("source_file") or "—"),
                f"{float(rec.get('confidence', 0)):.2f}",
                str(rec.get("status", "")),
            ]
        )
    return rows


def _add_provenance_table(doc: Document, records: list[dict[str, Any]]) -> None:
    _add_heading(doc, "Provenance Appendix", level=1)
    headers = [
        "Metric",
        "Period",
        "Approved Value",
        "Original Unit",
        "Page",
        "Source Document",
        "Source File",
        "Confidence",
        "Status",
    ]
    rows = _provenance_rows(records)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(rows, start=1):
        for ci, cell in enumerate(row):
            table.rows[ri].cells[ci].text = cell


def _add_cio_box(doc: Document) -> None:
    _add_heading(doc, "CIO / Fund Manager", level=1)
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = (
        "Investment view and recommendation to be completed by the fund manager.\n\n"
        "Rating: _______________\n"
        "Outlook: _______________\n"
        "Key risks: _______________"
    )
    # Light border emphasis
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:color"), "003366")
        borders.append(el)
    tc_pr.append(borders)


def generate_credit_review_report(
    *,
    reviewed_records: list[dict[str, Any]],
    commentary: dict[str, Any],
    warnings: list[str],
    output_path: Path,
    issuer_name: str | None = None,
    on_status: Callable[[str], None] | None = None,
) -> Path:
    """Build credit_review_report.docx at output_path."""

    def _status(msg: str) -> None:
        if on_status:
            on_status(msg)

    _status("Building analytical credit review report…")
    issuer = issuer_name or _issuer_name_from_records(reviewed_records)
    table1, table2 = split_records_by_table(reviewed_records)
    t1_periods = periods_from_records(table1) or TABLE1_PERIODS
    t2_periods = periods_from_records(table2) or TABLE2_PERIODS
    yearly_df = pivot_review_table(table1, t1_periods)
    half_df = pivot_review_table(table2, t2_periods)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE_PT)

    # --- Title page ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Credit Review Report")
    _set_run_font(tr, bold=True, size_pt=18)
    tr.font.color.rgb = HEADING_COLOR

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(issuer)
    _set_run_font(sr, bold=True, size_pt=14)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(datetime.now().strftime("%d %B %Y"))
    _set_run_font(dr)

    doc.add_page_break()

    # --- Issuer overview ---
    _add_heading(doc, "Issuer Overview", level=1)
    _add_body(
        doc,
        f"This credit review report for {issuer} is prepared from disclosed standalone "
        "annual report and investor presentation data. All monetary figures are in "
        "₹ crore unless stated otherwise. Ratios are expressed in percent. "
        "Values marked as not disclosed were not explicitly found in source documents.",
    )
    # --- Yearly table ---
    doc.add_page_break()
    _status("Formatting yearly financial tables…")
    _add_table_from_dataframe(
        doc,
        yearly_df,
        "Yearly Financials (March year-end, standalone)",
    )

    # --- Half-year table ---
    _status("Formatting half-year financial tables…")
    _add_table_from_dataframe(
        doc,
        half_df,
        "Half-Year Financials (H1FY26 / H1FY25, investor presentation)",
    )

    # --- Commentary (institutional sections — not per-metric lines) ---
    doc.add_page_break()
    _status("Injecting approved commentary…")
    _add_heading(doc, "Commentary", level=1)
    sections = commentary.get("sections") or []
    if not sections:
        for para in commentary.get("paragraphs", []):
            _add_body(doc, para)
    else:
        for section in sections:
            _add_heading(doc, section.get("title", "Section"), level=2)
            _add_body(doc, section.get("paragraph", ""))

    # --- Validation notes ---
    doc.add_page_break()
    _status("Adding validation notes and provenance…")
    _add_heading(doc, "Validation Notes", level=1)
    if warnings:
        for w in warnings:
            _add_body(doc, f"• {w}")
    else:
        _add_body(doc, "No validation warnings on approved values.")

    # --- Provenance ---
    doc.add_page_break()
    _add_provenance_table(doc, reviewed_records)

    # --- CIO box ---
    doc.add_page_break()
    _add_cio_box(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    _status("Preparing DOCX export…")
    doc.save(str(output_path))
    logger.info("Report saved to %s", output_path)
    return output_path
