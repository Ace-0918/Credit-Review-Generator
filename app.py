"""
Credit Review Report Generator — Phases 1–5.

Phase 1: PDF upload + page-wise extraction + document classification.
Phase 2: Deterministic financial metric extraction into two tables.
Phase 3: Human-in-the-loop review, manual edit, and approval workflow.
Phase 4: Rule-based commentary from approved values.
Phase 5: Analytical DOCX report generation.
Phase 6: Enterprise template formatter (optional uploaded .docx).
"""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Any

import pandas as pd
import streamlit as st

# Make `services/`, `utils/`, `data/` importable when launched via Streamlit
PROJECT_ROOT = Path(__file__).resolve().parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from services.document_classifier import classify_document, display_label  # noqa: E402
from services.extractor import run_financial_extraction  # noqa: E402
from services.pdf_reader import (  # noqa: E402
    extract_pages_from_pdf,
    has_no_extracted_text,
    is_likely_scanned_pdf,
    preview_tables_in_pdf,
    save_uploaded_pdf,
    total_table_count,
)
from services.commentary_generator import (  # noqa: E402
    generate_commentary,
    save_commentary_json,
)
from services.report_generator import generate_credit_review_report  # noqa: E402
from services.template_formatter import (  # noqa: E402
    FINAL_DOCX_NAME,
    FINAL_PDF_NAME,
    format_enterprise_report,
    resolve_template_path,
)
from services.review_manager import (  # noqa: E402
    COL_APPROVED,
    COL_CONFIDENCE,
    COL_EXTRACTED,
    COL_MANUAL_EDIT,
    COL_METRIC,
    COL_NOTES,
    COL_ORIGINAL_UNIT,
    COL_PAGE,
    COL_PERIOD,
    COL_SOURCE_DOC,
    COL_SOURCE_FILE,
    COL_STATUS,
    COL_VALUE_CRORE,
    apply_status_to_records,
    build_review_records,
    dataframe_to_records,
    merge_table_records,
    pivot_review_table,
    records_to_csv_bytes,
    records_to_dataframe,
    revalidate_approved,
    split_records_by_table,
)
from utils.live_status import LiveStatus  # noqa: E402
from utils.constants import (  # noqa: E402
    ALLOWED_EXTENSIONS,
    DATA_DIR,
    UPLOADS_DIR,
    DOC_TYPE_ANNUAL_REPORT,
    DOC_TYPE_CONCALL_TRANSCRIPT,
    DOC_TYPE_INVESTOR_PRESENTATION,
    OUTPUT_DIR,
)
from utils.logger import setup_logger  # noqa: E402

logger = setup_logger()

# -----------------------------------------------------------------------------
# Session state keys (kept stable across reruns)
# -----------------------------------------------------------------------------
SS_RAW = "raw_extraction_records"
SS_REVIEWED = "reviewed_extraction_records"
SS_APPROVED = "extraction_approved"
SS_WARNINGS = "review_warnings"
SS_PHASE1 = "phase1_results"
SS_SOURCE_MAP = "source_map"
SS_EDITOR_VERSION = "review_editor_version"
SS_COMMENTARY = "commentary_payload"
SS_COMMENTARY_DONE = "commentary_done"
SS_REPORT_DONE = "report_done"
SS_REPORT_PATH = "report_docx_path"
SS_TEMPLATE_BYTES = "enterprise_template_bytes"
SS_TEMPLATE_NAME = "enterprise_template_name"
SS_PHASE6_DONE = "phase6_done"
SS_FINAL_DOCX = "final_docx_path"
SS_FINAL_PDF = "final_pdf_path"

COMMENTARY_JSON_NAME = "commentary.json"
REPORT_DOCX_NAME = "credit_review_report.docx"
TEMPLATES_DIR = PROJECT_ROOT / "templates"


# -----------------------------------------------------------------------------
# Phase 1 helpers (unchanged from earlier phases)
# -----------------------------------------------------------------------------
def is_pdf_file(uploaded_file: Any) -> bool:
    if uploaded_file is None:
        return False
    name = getattr(uploaded_file, "name", "") or ""
    return Path(name).suffix.lower() in ALLOWED_EXTENSIONS


def normalize_upload_list(uploaded: Any) -> list[Any]:
    if uploaded is None:
        return []
    files = uploaded if isinstance(uploaded, list) else [uploaded]
    return [f for f in files if is_pdf_file(f)]


def read_upload_bytes(uploaded_file: Any) -> bytes:
    uploaded_file.seek(0)
    return uploaded_file.getvalue()


def run_extraction_for_file(
    uploaded_file: Any,
    upload_slot: str,
    category_label: str,
) -> dict[str, Any] | None:
    filename = uploaded_file.name
    pdf_bytes = read_upload_bytes(uploaded_file)
    try:
        saved_path = save_uploaded_pdf(pdf_bytes, filename, UPLOADS_DIR)
        pages = extract_pages_from_pdf(pdf_bytes, filename)
        doc_type = classify_document(filename, upload_slot=upload_slot)
        table_preview = preview_tables_in_pdf(pdf_bytes, filename)
        tables_found = total_table_count(table_preview)
    except ValueError as exc:
        logger.error("Extraction error for %s: %s", filename, exc)
        st.error(f"**{category_label}** — `{filename}`: {exc}")
        return None
    except Exception as exc:
        logger.exception("Unexpected error for %s", filename)
        st.error(
            f"**{category_label}** — `{filename}`: An unexpected error occurred — {exc}"
        )
        return None

    return {
        "category_label": category_label,
        "label": f"{category_label} — {filename}",
        "filename": filename,
        "pdf_bytes": pdf_bytes,
        "saved_path": str(saved_path),
        "pages": pages,
        "page_count": len(pages),
        "doc_type": doc_type,
        "doc_type_display": display_label(doc_type),
        "table_preview": table_preview,
        "tables_found": tables_found,
        "no_text": has_no_extracted_text(pages),
        "likely_scanned": is_likely_scanned_pdf(pages),
    }


def build_extraction_queue(
    annual_files: list[Any],
    investor_files: list[Any],
    concall_files: list[Any],
) -> list[tuple[Any, str, str]]:
    queue: list[tuple[Any, str, str]] = []
    for f in annual_files:
        queue.append((f, DOC_TYPE_ANNUAL_REPORT, "Annual Report"))
    for f in investor_files:
        queue.append((f, DOC_TYPE_INVESTOR_PRESENTATION, "Investor Presentation"))
    for f in concall_files:
        queue.append((f, DOC_TYPE_CONCALL_TRANSCRIPT, "Concall Transcript"))
    return queue


def _scan_status_message(upload_slot: str, filename: str) -> str:
    """Contextual Phase 1 status text by document type."""
    if upload_slot == DOC_TYPE_ANNUAL_REPORT:
        return f"Reading uploaded annual report: {filename}…"
    if upload_slot == DOC_TYPE_INVESTOR_PRESENTATION:
        return f"Reading investor presentation: {filename}…"
    if upload_slot == DOC_TYPE_CONCALL_TRANSCRIPT:
        return f"Scanning concall transcript: {filename}…"
    return f"Scanning document: {filename}…"


def run_batch_extraction(
    queue: list[tuple[Any, str, str]],
    *,
    live: LiveStatus | None = None,
) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    total = len(queue)
    if total == 0:
        return results

    if live is None:
        live = LiveStatus("Document scan")

    live.update("Preparing PDF scan…")
    for index, (uploaded_file, upload_slot, category_label) in enumerate(queue):
        filename = uploaded_file.name
        step = index + 1
        live.set_progress(
            (step - 1) / total,
            text=f"File {step} of {total}",
        )
        live.update(_scan_status_message(upload_slot, filename))
        if upload_slot == DOC_TYPE_ANNUAL_REPORT:
            live.update("Scanning financial sections…")
        elif upload_slot == DOC_TYPE_INVESTOR_PRESENTATION:
            live.update("Indexing presentation pages and tables…")

        result = run_extraction_for_file(
            uploaded_file, upload_slot=upload_slot, category_label=category_label
        )
        if result:
            results.append(result)

        live.set_progress(step / total, text=f"Completed {step} of {total}")

    live.clear_progress()
    live.update(f"Finished scanning {total} PDF(s).")
    return results


# -----------------------------------------------------------------------------
# Phase 1 / 2 driver — runs once, then caches in session_state
# -----------------------------------------------------------------------------
def run_full_pipeline(
    annual_files: list[Any],
    investor_files: list[Any],
    concall_files: list[Any],
) -> bool:
    """Execute Phase 1 + Phase 2 and store all artifacts in session_state."""
    # Clear previous extraction results
    for key in [SS_REVIEWED, SS_APPROVED, "llm_commentary"]:
        if key in st.session_state:
            del st.session_state[key]

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    UPLOADS_DIR.mkdir(parents=True, exist_ok=True)

    queue = build_extraction_queue(annual_files, investor_files, concall_files)
    logger.info("Starting Phase 1 — %d file(s)", len(queue))

    live = LiveStatus("Extraction")
    live.update("Starting document scan…")
    phase1_results = run_batch_extraction(queue, live=live)

    if not phase1_results:
        live.error("PDF scan failed. See output/extraction.log.")
        st.error("PDF scan failed. See `output/extraction.log`.")
        return False

    st.session_state[SS_PHASE1] = phase1_results

    gemini_key = st.session_state.get("gemini_api_key", "")
    if not gemini_key:
        try:
            gemini_key = st.secrets.get("GEMINI_API_KEY", "")
        except Exception:
            pass

    for res in phase1_results:
        res["vision_api_key"] = gemini_key or ""

    financial = run_financial_extraction(
        phase1_results,
        on_status=live.callback(),
        source_map=st.session_state.get(SS_SOURCE_MAP),
        fy_year=st.session_state.get("fy_year", 2026),
        year_end_month=st.session_state.get("year_end_month", "March"),
        h1_fy_year=st.session_state.get("h1_fy_year", st.session_state.get("fy_year", 2026)),
    )

    raw_records = merge_table_records(
        financial.table1_records, financial.table2_records
    )
    st.session_state[SS_RAW] = raw_records
    st.session_state["table1_records"] = financial.table1_records
    st.session_state["table2_records"] = financial.table2_records
    st.session_state[SS_REVIEWED] = build_review_records(raw_records)
    st.session_state[SS_APPROVED] = False
    st.session_state[SS_WARNINGS] = []
    st.session_state[SS_COMMENTARY] = None
    st.session_state[SS_COMMENTARY_DONE] = False
    st.session_state[SS_REPORT_DONE] = False
    st.session_state[SS_REPORT_PATH] = None
    st.session_state[SS_TEMPLATE_BYTES] = None
    st.session_state[SS_TEMPLATE_NAME] = None
    st.session_state[SS_PHASE6_DONE] = False
    st.session_state[SS_FINAL_DOCX] = None
    st.session_state[SS_FINAL_PDF] = None
    st.session_state[SS_EDITOR_VERSION] = st.session_state.get(SS_EDITOR_VERSION, 0) + 1

    # Persist a copy of the validation summary to disk (audit trail)
    summary_path = OUTPUT_DIR / "validation_summary.txt"
    summary_path.write_text(financial.validation_summary, encoding="utf-8")

    # ── Process any uploaded screenshots ─────────────────────────────────
    groq_key = st.session_state.get("groq_api_key", "")
    ann_shots = st.session_state.get("annual_screenshots_data", [])
    inv_shots = st.session_state.get("investor_screenshots_data", [])
    all_shots = list(ann_shots) + list(inv_shots)

    if all_shots and groq_key:
        from services.screenshot_extractor import (
            extract_from_screenshot,
            screenshot_to_review_records,
        )
        from data.metric_aliases import get_table1_periods, get_table2_periods
        _fy = st.session_state.get("fy_year", 2026)
        _month = st.session_state.get("year_end_month", "March")
        all_periods = (
            get_table1_periods(_fy, _month) +
            get_table2_periods(_fy, _month)
        )
        live.update("Extracting values from uploaded screenshots…")
        reviewed = st.session_state.get(SS_REVIEWED, [])
        total_found = 0
        for img_file in all_shots:
            img_bytes = img_file.read()
            ext = img_file.name.split(".")[-1].lower()
            fmt = "jpeg" if ext in ("jpg", "jpeg") else ext
            result = extract_from_screenshot(img_bytes, groq_key, fmt)
            if not result:
                continue
            new_records = screenshot_to_review_records(
                result,
                source_file=img_file.name,
                allowed_periods=all_periods,
            )
            by_key = {
                (r["metric"], r["period"]): i
                for i, r in enumerate(reviewed)
            }
            for new_rec in new_records:
                key = (new_rec["metric"], new_rec["period"])
                if key in by_key:
                    idx = by_key[key]
                    if reviewed[idx].get("approved_value") is None:
                        reviewed[idx].update(new_rec)
                        total_found += 1
                else:
                    reviewed.append(new_rec)
                    total_found += 1
        if total_found > 0:
            st.session_state[SS_REVIEWED] = reviewed
            live.update(f"Screenshots added {total_found} values.")
    # ── end screenshot processing ─────────────────────────────────────────

    live.success("Extraction complete — review and approve tables below.")

    from services.llm_extractor import get_last_gemini_error

    gemini_err = get_last_gemini_error()
    if gemini_err and (
        "RESOURCE_EXHAUSTED" in gemini_err
        or "429" in gemini_err
        or "quota" in gemini_err.lower()
    ):
        st.warning(
            "**Gemini API quota exceeded** — only pdfplumber values were extracted "
            "(FY24/25 and ratios often missing). Wait 1 hour or until tomorrow, then re-run. "
            "Check usage at [ai.dev/rate-limit](https://ai.dev/rate-limit). "
            "Details are in `output/extraction.log`."
        )
    elif gemini_err:
        st.warning(f"Gemini extraction failed: {gemini_err[:240]}")
    return True


def reset_pipeline() -> None:
    """Forget everything — return the app to the upload screen."""
    for key in (
        SS_RAW,
        SS_REVIEWED,
        SS_APPROVED,
        SS_WARNINGS,
        SS_PHASE1,
        SS_COMMENTARY,
        SS_COMMENTARY_DONE,
        SS_REPORT_DONE,
        SS_REPORT_PATH,
        SS_TEMPLATE_BYTES,
        SS_TEMPLATE_NAME,
        SS_PHASE6_DONE,
        SS_FINAL_DOCX,
        SS_FINAL_PDF,
    ):
        if key in st.session_state:
            del st.session_state[key]
    st.session_state[SS_EDITOR_VERSION] = (
        st.session_state.get(SS_EDITOR_VERSION, 0) + 1
    )


def render_workflow_progress() -> None:
    """Show pipeline phase status across Phases 1–6."""
    p1 = SS_PHASE1 in st.session_state and bool(st.session_state.get(SS_PHASE1))
    p2 = SS_RAW in st.session_state and bool(st.session_state.get(SS_RAW))
    p3 = bool(st.session_state.get(SS_APPROVED))
    p4 = bool(st.session_state.get(SS_COMMENTARY_DONE))
    p5 = bool(st.session_state.get(SS_REPORT_DONE))
    p6 = bool(st.session_state.get(SS_PHASE6_DONE))

    def mark(done: bool) -> str:
        return "✅" if done else "⬜"

    st.markdown(
        f"{mark(p1)} **Phase 1 — Scan** &nbsp;&nbsp;|&nbsp;&nbsp; "
        f"{mark(p2)} **Phase 2 — Extraction** &nbsp;&nbsp;|&nbsp;&nbsp; "
        f"{mark(p3)} **Phase 3 — Review** &nbsp;&nbsp;|&nbsp;&nbsp; "
        f"{mark(p4)} **Phase 4 — Commentary** &nbsp;&nbsp;|&nbsp;&nbsp; "
        f"{mark(p5)} **Phase 5 — Analytical Report** &nbsp;&nbsp;|&nbsp;&nbsp; "
        f"{mark(p6)} **Phase 6 — Enterprise Formatter**"
    )
    st.divider()


# -----------------------------------------------------------------------------
# Phase 1 scan summary (collapsed inside the review screen)
# -----------------------------------------------------------------------------
def render_phase1_scan_summary(phase1_results: list[dict[str, Any]]) -> None:
    with st.expander("Document scan summary (Phase 1)", expanded=False):
        for result in phase1_results:
            st.markdown(
                f"**{result['category_label']}** — `{result['filename']}` · "
                f"{result['page_count']} pages · {result['tables_found']} tables · "
                f"type `{result['doc_type']}`"
            )
            if result["no_text"]:
                st.warning(f"{result['filename']}: No text extracted.")
            elif result["likely_scanned"]:
                st.warning(f"{result['filename']}: Possible scanned PDF.")
            if result["tables_found"] == 0:
                st.warning(f"{result['filename']}: No tables detected.")


# -----------------------------------------------------------------------------
# Phase 3 — Review UI
# -----------------------------------------------------------------------------
def _editor_column_config() -> dict[str, Any]:
    """Column visibility / edit rules for the review data_editor."""
    return {
        COL_METRIC: st.column_config.TextColumn("Metric", disabled=True),
        COL_PERIOD: st.column_config.TextColumn("Period", disabled=True),
        COL_EXTRACTED: st.column_config.NumberColumn(
            "Extracted Value",
            disabled=True,
            format="%.2f",
            help="Value as printed in the PDF (in the original unit).",
        ),
        COL_APPROVED: st.column_config.NumberColumn(
            "Approved Value (Crore / %)",
            format="%.2f",
            help=(
                "Editable. Final normalized value in ₹ crore or percent. "
                "Converted Value (Crore) stays in sync when you save."
            ),
        ),
        COL_ORIGINAL_UNIT: st.column_config.TextColumn("Original Unit", disabled=True),
        COL_VALUE_CRORE: st.column_config.NumberColumn(
            "Converted Value (Crore)",
            disabled=True,
            format="%.2f",
        ),
        COL_SOURCE_DOC: st.column_config.TextColumn("Source Document", disabled=True),
        COL_SOURCE_FILE: st.column_config.TextColumn("Source File", disabled=True),
        COL_PAGE: st.column_config.TextColumn("Page Number", disabled=True),
        COL_CONFIDENCE: st.column_config.NumberColumn(
            "Confidence",
            disabled=True,
            format="%.2f",
        ),
        COL_STATUS: st.column_config.TextColumn("Status", disabled=True),
        COL_NOTES: st.column_config.TextColumn(
            "Notes",
            help="Optional review notes.",
        ),
        COL_MANUAL_EDIT: st.column_config.CheckboxColumn(
            COL_MANUAL_EDIT,
            disabled=True,
            help="Checked automatically when Approved Value differs from the extraction baseline.",
        ),
    }


def _render_review_editor(
    section_title: str,
    section_subtitle: str,
    table_records: list[dict[str, Any]],
    periods: tuple[str, ...],
    editor_key: str,
) -> pd.DataFrame:
    """Render one of the two review sections (yearly or half-year)."""
    st.subheader(section_title)
    if section_subtitle:
        st.caption(section_subtitle)

    pivot_df = pivot_review_table(table_records, periods)
    st.markdown("**Approved values (current snapshot)**")
    st.dataframe(pivot_df, use_container_width=True, hide_index=True)

    st.markdown("**Detailed provenance — edit Approved Value if needed**")
    editor_df = records_to_dataframe(table_records)

    edited_df = st.data_editor(
        editor_df,
        key=editor_key,
        use_container_width=True,
        hide_index=True,
        num_rows="fixed",
        column_config=_editor_column_config(),
        column_order=[
            COL_METRIC,
            COL_PERIOD,
            COL_EXTRACTED,
            COL_APPROVED,
            COL_ORIGINAL_UNIT,
            COL_VALUE_CRORE,
            COL_SOURCE_DOC,
            COL_PAGE,
            COL_CONFIDENCE,
            COL_STATUS,
            COL_MANUAL_EDIT,
            COL_NOTES,
            COL_SOURCE_FILE,
        ],
    )

    # ── Screenshot upload for missing values ──────────────────────────────
    missing_metrics = [
        rec["metric"] for rec in table_records
        if rec.get("approved_value") is None
    ]
    if missing_metrics:
        with st.expander(
            f"📷 Upload screenshots for missing values ({len(missing_metrics)} missing)",
            expanded=False,
        ):
            st.caption(
                "Upload a photo or screenshot of the exact page in the PDF "
                "that contains the missing values. Groq Vision will extract "
                "the numbers directly from the image."
            )
            screenshot_files = st.file_uploader(
                "Upload page screenshots (JPG, PNG)",
                type=["jpg", "jpeg", "png", "webp"],
                accept_multiple_files=True,
                key=f"screenshots_{editor_key}",
            )
            groq_key = st.session_state.get("groq_api_key", "")
            if screenshot_files and groq_key:
                if st.button(
                    "Extract from screenshots",
                    key=f"btn_screenshot_{editor_key}",
                ):
                    from services.screenshot_extractor import (
                        extract_from_screenshot,
                        screenshot_to_review_records,
                    )
                    from data.metric_aliases import get_table1_periods, get_table2_periods
                    _fy = st.session_state.get("fy_year", 2026)
                    _month = st.session_state.get("year_end_month", "March")
                    all_periods = (
                        get_table1_periods(_fy, _month) +
                        get_table2_periods(_fy, _month)
                    )
                    total_found = 0
                    for img_file in screenshot_files:
                        img_bytes = img_file.read()
                        ext = img_file.name.split(".")[-1].lower()
                        fmt = "jpeg" if ext in ("jpg", "jpeg") else ext
                        with st.spinner(f"Reading {img_file.name}…"):
                            result = extract_from_screenshot(
                                img_bytes, groq_key, fmt
                            )
                        if not result:
                            st.warning(f"Could not extract from {img_file.name}")
                            continue
                        new_records = screenshot_to_review_records(
                            result,
                            source_file=img_file.name,
                            allowed_periods=all_periods,
                        )
                        # Merge into reviewed records
                        reviewed = st.session_state.get(SS_REVIEWED, [])
                        by_key = {
                            (r["metric"], r["period"]): i
                            for i, r in enumerate(reviewed)
                        }
                        for new_rec in new_records:
                            key = (new_rec["metric"], new_rec["period"])
                            if key in by_key:
                                idx = by_key[key]
                                if reviewed[idx].get("approved_value") is None:
                                    reviewed[idx].update(new_rec)
                                    total_found += 1
                        st.session_state[SS_REVIEWED] = reviewed
                    if total_found > 0:
                        st.success(f"Extracted {total_found} values from screenshots.")
                        st.rerun()
                    else:
                        st.info("No new values found in screenshots.")
            elif screenshot_files and not groq_key:
                st.warning("Enter your Groq API key above to extract from screenshots.")
    # ── end screenshot upload ─────────────────────────────────────────────

    return edited_df


def render_review_workflow() -> None:
    """Top-level Phase 3 review and approval screen."""
    if SS_REVIEWED not in st.session_state:
        return

    from data.metric_aliases import get_table1_periods, get_table2_periods
    _fy = st.session_state.get("fy_year", 2026)
    _month = st.session_state.get("year_end_month", "March")
    TABLE1_PERIODS = get_table1_periods(_fy, _month)
    TABLE2_PERIODS = get_table2_periods(_fy, _month)

    reviewed: list[dict[str, Any]] = st.session_state[SS_REVIEWED]
    table1, table2 = split_records_by_table(reviewed)

    st.title("Review Extracted Financial Data")
    st.caption("Please verify the extracted values before generating the report.")

    render_workflow_progress()
    render_phase1_scan_summary(st.session_state.get(SS_PHASE1, []))

    version = st.session_state.get(SS_EDITOR_VERSION, 0)

    edited_t1 = _render_review_editor(
        section_title="Yearly Financials",
        section_subtitle=(
            "From annual report or investor presentation — standalone, March year-end. "
            f"Periods: {', '.join(TABLE1_PERIODS)}."
        ),
        table_records=table1,
        periods=TABLE1_PERIODS,
        editor_key=f"editor_table1_v{version}",
    )

    edited_t2 = _render_review_editor(
        section_title="Half-Year Financials",
        section_subtitle=(
            "From investor presentation — half-year September ended only "
            "(H1FY26, H1FY25). Q1/Q2 single quarters are not used."
        ),
        table_records=table2,
        periods=TABLE2_PERIODS,
        editor_key=f"editor_table2_v{version}",
    )

    if st.button("🔍 View Raw Extraction Data", key="btn_raw_debug"):
        st.session_state["show_raw_debug"] = True

    if st.session_state.get("show_raw_debug"):
        import pandas as pd
        import io

        all_records = (
            st.session_state.get("table1_records", []) +
            st.session_state.get("table2_records", [])
        )

        raw_rows = []
        for rec in all_records:
            val_orig = rec.get("value_original")
            val_crore = rec.get("value_crore")
            unit = rec.get("unit") or "unknown"
            raw = rec.get("raw_text") or "—"

            # Build conversion explanation
            if val_orig is not None and val_crore is not None:
                try:
                    if unit == "thousand":
                        conv = f"{val_orig} ÷ 10,000 = {val_crore}"
                    elif unit == "lakh":
                        conv = f"{val_orig} ÷ 100 = {val_crore}"
                    elif unit == "crore":
                        conv = f"{val_orig} (direct) = {val_crore}"
                    elif unit == "million":
                        conv = f"{val_orig} × 0.1 = {val_crore}"
                    else:
                        conv = f"{val_orig} → {val_crore}"
                except Exception:
                    conv = "—"
            else:
                conv = "—"

            raw_rows.append({
                "Period":            rec.get("period", ""),
                "Metric":            rec.get("metric", ""),
                "Raw Text from PDF": raw,
                "Unit Detected":     unit,
                "Conversion":        conv,
                "Final Value (cr/%)":rec.get("value_crore") or "Not Disclosed",
                "Page No.":          rec.get("page_number") or "—",
                "PDF Section":       rec.get("source_section") or "—",
                "Row Label in PDF":  rec.get("row_label") or "—",
                "Column Header":     rec.get("column_header") or "—",
                "Confidence":        f"{float(rec.get('confidence', 0)):.2f}",
                "Status":            rec.get("status", ""),
                "Source File":       rec.get("source_filename") or rec.get("source_file") or "—",
            })

        if raw_rows:
            raw_df = pd.DataFrame(raw_rows)

            # Show inline with filters
            col1, col2 = st.columns(2)
            with col1:
                filter_metric = st.multiselect(
                    "Filter by metric",
                    options=sorted(raw_df["Metric"].unique().tolist()),
                    default=[],
                    key="raw_filter_metric",
                )
            with col2:
                filter_status = st.multiselect(
                    "Filter by status",
                    options=["extracted", "missing", "Low Confidence", "Warning"],
                    default=[],
                    key="raw_filter_status",
                )

            filtered = raw_df.copy()
            if filter_metric:
                filtered = filtered[filtered["Metric"].isin(filter_metric)]
            if filter_status:
                filtered = filtered[
                    filtered["Status"].str.lower().isin(
                        [s.lower() for s in filter_status]
                    )
                ]

            st.dataframe(filtered, use_container_width=True, height=450)

            # Excel download
            buffer = io.BytesIO()
            with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
                # Sheet 1: All data
                raw_df.to_excel(
                    writer, index=False, sheet_name="All Extracted"
                )
                # Sheet 2: Missing only
                missing_df = raw_df[raw_df["Status"] == "missing"]
                if not missing_df.empty:
                    missing_df.to_excel(
                        writer, index=False, sheet_name="Missing Values"
                    )
                # Sheet 3: Warnings only
                warn_df = raw_df[raw_df["Status"].str.contains("Warning", case=False, na=False)]
                if not warn_df.empty:
                    warn_df.to_excel(
                        writer, index=False, sheet_name="Warnings"
                    )

                all_hits_rows = []
                for doc_ctx in st.session_state.get("doc_contexts", []):
                    for hit in getattr(doc_ctx, "all_extraction_hits", []):
                        all_hits_rows.append({
                            "Metric": hit.metric,
                            "Period": hit.period,
                            "Value (cr/%)": hit.value_crore,
                            "Page": hit.page_number,
                            "Row Label in PDF": hit.row_label,
                            "Column Header": hit.column_header,
                            "Confidence": f"{hit.confidence:.2f}",
                            "Section": hit.source_section,
                            "Source File": hit.source_file,
                        })

                if all_hits_rows:
                    all_hits_df = pd.DataFrame(all_hits_rows)
                    all_hits_df.to_excel(
                        writer, index=False, sheet_name="All Candidates"
                    )
            buffer.seek(0)

            st.download_button(
                label="📥 Download full raw extraction Excel",
                data=buffer,
                file_name="raw_extraction.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="dl_raw_excel",
            )

            if st.button("Close", key="btn_close_raw"):
                st.session_state["show_raw_debug"] = False
                st.rerun()
        else:
            st.info("Run extraction first.")

    _render_review_action_bar(edited_t1, edited_t2, table1, table2)
    _render_warnings_panel()
    _render_approval_banner()

    if st.session_state.get(SS_APPROVED):
        _render_phase4_commentary()
        _render_phase5_report()
        _render_phase6_formatter()


def _commit_edits(
    edited_t1: pd.DataFrame,
    edited_t2: pd.DataFrame,
    table1_base: list[dict[str, Any]],
    table2_base: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Merge editor edits back into review records and persist to session_state."""
    updated_t1 = dataframe_to_records(edited_t1, table1_base)
    updated_t2 = dataframe_to_records(edited_t2, table2_base)
    combined = merge_table_records(updated_t1, updated_t2)
    st.session_state[SS_REVIEWED] = combined
    return combined


def _render_review_action_bar(
    edited_t1: pd.DataFrame,
    edited_t2: pd.DataFrame,
    table1_base: list[dict[str, Any]],
    table2_base: list[dict[str, Any]],
) -> None:
    """Save / Approve / Reset / Download buttons."""
    st.divider()
    col_save, col_approve, col_reset, col_csv = st.columns([1, 1, 1, 1.2])

    if col_save.button("Save Reviewed Data", use_container_width=True):
        _commit_edits(edited_t1, edited_t2, table1_base, table2_base)
        st.session_state[SS_APPROVED] = False
        st.session_state[SS_WARNINGS] = revalidate_approved(
            st.session_state[SS_REVIEWED]
        )
        st.success("Reviewed data saved.")
        st.rerun()

    if col_approve.button(
        "Approve Extraction",
        type="primary",
        use_container_width=True,
    ):
        combined = _commit_edits(edited_t1, edited_t2, table1_base, table2_base)
        warnings = revalidate_approved(combined)
        apply_status_to_records(combined, final_approval=True)
        st.session_state[SS_REVIEWED] = combined
        st.session_state[SS_WARNINGS] = warnings
        st.session_state[SS_APPROVED] = True
        st.session_state[SS_COMMENTARY] = None
        st.session_state[SS_COMMENTARY_DONE] = False
        st.session_state[SS_REPORT_DONE] = False
        st.session_state[SS_REPORT_PATH] = None
        st.session_state[SS_PHASE6_DONE] = False
        st.session_state[SS_FINAL_DOCX] = None
        st.session_state[SS_FINAL_PDF] = None
        if warnings:
            st.warning(
                "Warning: Some values failed validation. Please verify before proceeding."
            )
        st.rerun()

    if col_reset.button("Reset Edits", use_container_width=True):
        raw_records = st.session_state.get(SS_RAW, [])
        st.session_state[SS_REVIEWED] = build_review_records(raw_records)
        st.session_state[SS_APPROVED] = False
        st.session_state[SS_WARNINGS] = []
        st.session_state[SS_EDITOR_VERSION] = (
            st.session_state.get(SS_EDITOR_VERSION, 0) + 1
        )
        st.info("Edits reset to the originally extracted values.")
        st.rerun()

    # Download always reflects the *saved* reviewed records
    csv_bytes = records_to_csv_bytes(st.session_state[SS_REVIEWED])
    col_csv.download_button(
        "Download Reviewed Extraction CSV",
        data=csv_bytes,
        file_name="reviewed_extraction.csv",
        mime="text/csv",
        use_container_width=True,
    )


def _render_warnings_panel() -> None:
    warnings: list[str] = st.session_state.get(SS_WARNINGS, [])
    if not warnings:
        return
    st.subheader("Validation Warnings")
    for w in warnings:
        st.warning(f"⚠ {w}")


def _render_approval_banner() -> None:
    if not st.session_state.get(SS_APPROVED):
        return
    reviewed: list[dict[str, Any]] = st.session_state.get(SS_REVIEWED, [])
    counts: dict[str, int] = {}
    for rec in reviewed:
        counts[rec["status"]] = counts.get(rec["status"], 0) + 1

    st.divider()
    st.success("Extraction approved. Proceed to **Generate Commentary** below.")
    summary_lines = [f"- {status}: {count}" for status, count in counts.items()]
    st.markdown("Final status counts:\n" + "\n".join(summary_lines))


def _render_phase4_commentary() -> None:
    """Phase 4 — deterministic commentary from approved values."""
    st.divider()
    st.header("Generate Commentary")
    st.caption(
        "Rule-based narrative from approved values only. "
        "Institutional section paragraphs (Business Profile, Profitability, Capitalisation, "
        "Liquidity). Missing metrics are noted briefly; no numbers are invented."
    )

    if not st.session_state.get(SS_COMMENTARY_DONE):
        if st.button("Generate Commentary", type="primary", key="btn_gen_commentary"):
            reviewed = st.session_state.get(SS_REVIEWED, [])
            live = LiveStatus("Commentary")
            payload = generate_commentary(reviewed, on_status=live.callback())
            out_path = OUTPUT_DIR / COMMENTARY_JSON_NAME
            save_commentary_json(payload, out_path)
            st.session_state[SS_COMMENTARY] = payload
            st.session_state[SS_COMMENTARY_DONE] = True
            st.session_state[SS_REPORT_DONE] = False
            st.session_state[SS_REPORT_PATH] = None
            st.session_state[SS_PHASE6_DONE] = False
            st.session_state[SS_FINAL_DOCX] = None
            st.session_state[SS_FINAL_PDF] = None
            live.success("Commentary ready for report generation.")
            st.success(f"Commentary saved to `{out_path}`.")
            st.rerun()
        return

    st.success("Commentary generated.")
    payload = st.session_state.get(SS_COMMENTARY) or {}
    sections = payload.get("sections", [])
    if sections:
        with st.expander("Preview — Institutional commentary", expanded=True):
            for section in sections:
                st.markdown(f"**{section.get('title', 'Section')}**")
                st.markdown(section.get("paragraph", ""))
    else:
        with st.expander("Preview — Commentary", expanded=True):
            for line in payload.get("paragraphs", []):
                st.markdown(line)

    json_path = OUTPUT_DIR / COMMENTARY_JSON_NAME
    if json_path.is_file():
        st.download_button(
            "Download commentary.json",
            data=json_path.read_bytes(),
            file_name=COMMENTARY_JSON_NAME,
            mime="application/json",
            key="dl_commentary_json",
        )

    # ── LLM Commentary (Groq) ─────────────────────────────────────────────
    st.markdown("---")
    st.subheader("AI Commentary (Groq)")

    # Load default key from secrets if available
    default_key = ""
    try:
        default_key = st.secrets.get("GROQ_API_KEY", "")
    except Exception:
        pass

    groq_key = st.text_input(
        "Groq API Key",
        value=default_key,
        type="password",
        placeholder="gsk_...",
        help="A default key is pre-loaded. You can replace it with your own free key from console.groq.com",
        key="groq_api_key_input",
    )
    if groq_key == default_key and default_key:
        st.caption("Using default API key. Replace with your own from console.groq.com if needed.")
    st.session_state["groq_api_key"] = groq_key

    # Gemini API key (also set on upload screen; keep in sync here for re-runs)
    default_gemini = ""
    try:
        default_gemini = st.secrets.get("GEMINI_API_KEY", "")
    except Exception:
        pass

    gemini_key = st.text_input(
        "Gemini API Key",
        value=default_gemini,
        type="password",
        placeholder="AQ.... or AIza...",
        help="Used during PDF extraction. pdfplumber fills any cells Gemini misses.",
        key="gemini_api_key_input",
    )
    if gemini_key == default_gemini and default_gemini:
        st.caption(
            "Using default Gemini key. "
            "Replace with your own from aistudio.google.com if needed."
        )
    st.session_state["gemini_api_key"] = gemini_key

    if groq_key and st.button("Generate AI Commentary", key="btn_groq_commentary"):
        from services.llm_commentary import generate_llm_commentary
        from services.report_generator import _issuer_name_from_records

        live = LiveStatus("AI Commentary")
        try:
            issuer = _issuer_name_from_records(
                st.session_state.get(SS_REVIEWED, [])
            )
            llm_sections = generate_llm_commentary(
                reviewed_records=st.session_state.get(SS_REVIEWED, []),
                issuer_name=issuer,
                api_key=groq_key,
                on_status=live.callback(),
            )
            st.session_state["llm_commentary"] = llm_sections
            live.success("AI commentary generated successfully.")
        except Exception as e:
            live.error(f"Commentary generation failed: {e}")

    # Display generated sections if available
    if st.session_state.get("llm_commentary"):
        llm_sections = st.session_state["llm_commentary"]
        section_labels = {
            "company_profile":  "Company Profile",
            "profitability":    "Profitability",
            "asset_quality":    "Asset Quality",
            "capitalisation":   "Capitalisation",
            "liquidity":        "Liquidity",
            "recommendation":   "Recommendation",
        }
        for key, label in section_labels.items():
            if key in llm_sections:
                st.markdown(f"**{label}**")
                st.write(llm_sections[key])
                st.markdown("")
    # ──────────────────────────────────────────────────────────────────────

    if st.button("Regenerate Commentary", key="btn_regen_commentary"):
        st.session_state[SS_COMMENTARY_DONE] = False
        st.session_state[SS_COMMENTARY] = None
        st.session_state[SS_REPORT_DONE] = False
        st.session_state[SS_REPORT_PATH] = None
        st.session_state[SS_PHASE6_DONE] = False
        st.session_state[SS_FINAL_DOCX] = None
        st.session_state[SS_FINAL_PDF] = None
        st.rerun()


def _render_phase5_report() -> None:
    """Phase 5 — analytical DOCX report generation."""
    if not st.session_state.get(SS_COMMENTARY_DONE):
        st.info("Complete **Generate Commentary** (Phase 4) before generating the report.")
        return

    st.divider()
    st.header("Generate Analytical Report (Phase 5)")
    st.caption(
        "Internal analytical credit review (Word). Uses approved tables, commentary, "
        "and validation notes. Phase 6 can reformat this into your enterprise template."
    )

    report_path = OUTPUT_DIR / REPORT_DOCX_NAME

    if not st.session_state.get(SS_REPORT_DONE):
        if st.button("Generate Report", type="primary", key="btn_gen_report"):
            reviewed = st.session_state.get(SS_REVIEWED, [])
            commentary = st.session_state.get(SS_COMMENTARY) or {}
            warnings = st.session_state.get(SS_WARNINGS, [])
            live = LiveStatus("Report generation")
            generate_credit_review_report(
                reviewed_records=reviewed,
                commentary=commentary,
                warnings=warnings,
                output_path=report_path,
                on_status=live.callback(),
            )
            st.session_state[SS_REPORT_PATH] = str(report_path)
            st.session_state[SS_REPORT_DONE] = True
            st.session_state[SS_PHASE6_DONE] = False
            st.session_state[SS_FINAL_DOCX] = None
            st.session_state[SS_FINAL_PDF] = None
            live.success("Analytical credit review report ready.")
            st.success(f"Report saved to `{report_path}`.")
            st.rerun()
        return

    st.success("Report generated.")
    saved = st.session_state.get(SS_REPORT_PATH) or str(report_path)
    path = Path(saved)
    if path.is_file():
        st.download_button(
            "Download credit_review_report.docx",
            data=path.read_bytes(),
            file_name=REPORT_DOCX_NAME,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_report_docx",
        )
    if st.button("Regenerate Report", key="btn_regen_report"):
        st.session_state[SS_REPORT_DONE] = False
        st.session_state[SS_REPORT_PATH] = None
        st.session_state[SS_PHASE6_DONE] = False
        st.session_state[SS_FINAL_DOCX] = None
        st.session_state[SS_FINAL_PDF] = None
        st.rerun()


def _render_phase6_formatter() -> None:
    """Phase 6 — inject content into enterprise .docx template (uploaded or default)."""
    if not st.session_state.get(SS_REPORT_DONE):
        st.info("Complete **Generate Analytical Report** (Phase 5) before enterprise formatting.")
        return

    st.divider()
    st.header("Enterprise Formatter (Phase 6)")
    st.caption(
        "Replace all dynamic content in your enterprise .docx template while preserving "
        "fonts, borders, colours, and layout. Template text is treated as placeholder "
        "only — financial tables and commentary are overwritten with approved extraction."
    )

    uploaded = st.file_uploader(
        "Upload Enterprise Format Template (Optional)",
        type=["docx"],
        key="enterprise_template_upload",
        help="Fund-specific credit memo / Kotak / Aditya Birla format. Leave empty for default template.",
    )
    if uploaded is not None:
        st.session_state[SS_TEMPLATE_BYTES] = read_upload_bytes(uploaded)
        st.session_state[SS_TEMPLATE_NAME] = uploaded.name
        st.caption(f"Template loaded in session: `{uploaded.name}`")

    with st.expander("Supported placeholders (Mode B)", expanded=False):
        st.markdown(
            """
| Placeholder | Content |
|---|---|
| `{{COMPANY_NAME}}` / `{{ISSUER_NAME}}` | Issuer name |
| `{{DATE}}` / `{{REPORT_DATE}}` | Report date |
| `{{ISSUER_OVERVIEW}}` | Company profile text |
| `{{YEARLY_TABLE}}` | Yearly financial table |
| `{{HALFYEAR_TABLE}}` / `{{HALF_YEAR_TABLE}}` | Half-year table |
| `{{COMMENTARY}}` | Full commentary |
| `{{COMMENTARY_YEARLY}}` / `{{COMMENTARY_HALFYEAR}}` | Section commentary |
| `{{VALIDATION_NOTES}}` | Validation warnings |
| `{{RECOMMENDATION}}` / `{{CIO_FUND_MANAGER}}` | CIO / fund manager block |
            """
        )
        st.markdown(
            "**Smart section detection:** headings such as *Yearly Financials*, "
            "*Half-Year Financials*, *Profitability*, *Capitalisation*, *Liquidity*, "
            "*Commentary*, and *CIO / Fund Manager* have their **body text replaced** "
            "(heading style preserved). Financial tables are overwritten in-place by "
            "metric and period — stale template numbers are never kept."
        )

    if not st.session_state.get(SS_PHASE6_DONE):
        if st.button("Generate Enterprise Report", type="primary", key="btn_phase6"):
            reviewed = st.session_state.get(SS_REVIEWED, [])
            commentary = st.session_state.get(SS_COMMENTARY) or {}
            warnings = st.session_state.get(SS_WARNINGS, [])
            TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
            template_path = resolve_template_path(
                TEMPLATES_DIR,
                st.session_state.get(SS_TEMPLATE_BYTES),
                st.session_state.get(SS_TEMPLATE_NAME),
            )
            live = LiveStatus("Enterprise formatter")
            result = format_enterprise_report(
                template_path=template_path,
                reviewed_records=reviewed,
                commentary=commentary,
                warnings=warnings,
                output_dir=OUTPUT_DIR,
                on_status=live.callback(),
                llm_sections=st.session_state.get("llm_commentary"),
            )
            st.session_state[SS_FINAL_DOCX] = result["docx_path"]
            st.session_state[SS_FINAL_PDF] = result.get("pdf_path")
            st.session_state[SS_PHASE6_DONE] = True
            live.success("Credit review report ready.")
            st.success(f"Enterprise report saved: `{result['docx_path']}`")
            if not result.get("pdf_exported"):
                st.warning(
                    "PDF export was not available on this machine. "
                    "Install Microsoft Word or `pip install docx2pdf` for PDF output."
                )
            st.rerun()
        return

    st.success("Enterprise report generated.")
    docx_path = Path(st.session_state.get(SS_FINAL_DOCX) or OUTPUT_DIR / FINAL_DOCX_NAME)

    # ── Report Preview ─────────────────────────────────────────────────────
    st.markdown("---")
    st.subheader("Report Preview")

    with st.expander("Preview generated report before downloading", expanded=True):
        try:
            from docx import Document as _DocxDocument

            preview_doc = _DocxDocument(str(docx_path))

            # Show each section
            current_heading = None
            for para in preview_doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                # Detect headings by bold + larger font or heading style
                is_heading = (
                    para.style.name.startswith("Heading")
                    or any(
                        run.bold
                        and run.font.size
                        and run.font.size.pt >= 12
                        for run in para.runs
                        if run.text.strip()
                    )
                )
                if is_heading:
                    st.markdown(f"**{text}**")
                else:
                    st.write(text)

            # Show tables
            st.markdown("---")
            st.markdown("**Financial Tables**")
            for table_idx, table in enumerate(preview_doc.tables):
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                if rows:
                    # First row as header
                    try:
                        df = pd.DataFrame(rows[1:], columns=rows[0])
                        st.dataframe(df, use_container_width=True)
                    except Exception:
                        # Fallback if header row has issues
                        df = pd.DataFrame(rows)
                        st.dataframe(df, use_container_width=True)
                    if table_idx >= 3:
                        st.caption("(Further tables omitted from preview)")
                        break

        except Exception as e:
            st.warning(f"Preview could not be rendered: {e}")

    st.markdown("---")
    # ── Download buttons follow here (existing code unchanged) ──────────────

    if docx_path.is_file():
        st.download_button(
            f"Download {FINAL_DOCX_NAME}",
            data=docx_path.read_bytes(),
            file_name=FINAL_DOCX_NAME,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_final_docx",
        )
    pdf_path_str = st.session_state.get(SS_FINAL_PDF)
    if pdf_path_str:
        pdf_path = Path(pdf_path_str)
        if pdf_path.is_file():
            st.download_button(
                f"Download {FINAL_PDF_NAME}",
                data=pdf_path.read_bytes(),
                file_name=FINAL_PDF_NAME,
                mime="application/pdf",
                key="dl_final_pdf",
            )

    if st.button("Regenerate Enterprise Report", key="btn_regen_phase6"):
        st.session_state[SS_PHASE6_DONE] = False
        st.session_state[SS_FINAL_DOCX] = None
        st.session_state[SS_FINAL_PDF] = None
        st.rerun()


# -----------------------------------------------------------------------------
# Upload screen (shown only when extraction has not been run yet)
# -----------------------------------------------------------------------------
def render_upload_screen() -> None:
    st.title("Credit Review Report Generator")
    st.caption(
        "Upload financial PDFs, extract disclosed metrics, and validate before report generation."
    )

    st.header("Upload Documents")
    st.markdown(
        "PDF only. Multiple files per category allowed. "
        "**Yearly Financials** prefer annual report(s); if none are uploaded, "
        "the investor presentation is used. "
        "**Half-Year Financials** use investor presentation(s)."
    )

    col1, col2, col3 = st.columns(3)

    with col1:
        selected_fy = st.selectbox(
            "Yearly Financial Year",
            options=list(range(2020, 2031)),
            index=5,
            format_func=lambda y: f"FY{y}",
            key="selected_fy_year",
        )

    with col2:
        year_end_month = st.selectbox(
            "Year-end month",
            options=["March", "June", "September", "December"],
            index=0,
            key="selected_year_end_month",
        )

    with col3:
        selected_h1_fy = st.selectbox(
            "Half-Year Financial Year",
            options=list(range(2020, 2031)),
            index=6,
            format_func=lambda y: f"H1FY{str(y)[2:]}",
            key="selected_h1_fy_year",
            help="Select the half-year period from your investor PPT",
        )

    st.session_state["fy_year"] = selected_fy
    st.session_state["year_end_month"] = year_end_month
    st.session_state["h1_fy_year"] = selected_h1_fy

    from data.metric_aliases import get_table1_periods, get_table2_periods
    t1 = get_table1_periods(selected_fy, year_end_month)
    t2 = get_table2_periods(selected_h1_fy, year_end_month)

    col1, col2 = st.columns(2)
    with col1:
        st.info(
            "**Yearly periods:**\n\n"
            + "\n".join(f"• {p}" for p in t1)
        )
    with col2:
        st.info(
            "**Half-year periods:**\n\n"
            + "\n".join(f"• {p}" for p in t2)
        )

    col1, col2, col3 = st.columns(3)
    with col1:
        annual_uploads = st.file_uploader(
            "Annual Report PDFs (optional)",
            type=["pdf"],
            accept_multiple_files=True,
            key="annual_report",
        )
        annual_screenshots = st.file_uploader(
            "Or upload page screenshots (JPG/PNG)",
            type=["jpg", "jpeg", "png", "webp"],
            accept_multiple_files=True,
            key="annual_screenshots",
            help="Upload photos or screenshots of specific financial pages "
                 "if you don't have the full PDF.",
        )
    with col2:
        investor_uploads = st.file_uploader(
            "Investor Presentation PDFs (optional if annual report uploaded)",
            type=["pdf"],
            accept_multiple_files=True,
            key="investor_presentation",
        )
        investor_screenshots = st.file_uploader(
            "Or upload page screenshots (JPG/PNG)",
            type=["jpg", "jpeg", "png", "webp"],
            accept_multiple_files=True,
            key="investor_screenshots",
            help="Upload screenshots of H1 results pages if no full PPT available.",
        )
    with col3:
        concall_uploads = st.file_uploader(
            "Concall Transcript PDFs (optional)",
            type=["pdf"],
            accept_multiple_files=True,
            key="concall_transcript",
        )

    annual_files = normalize_upload_list(annual_uploads)
    investor_files = normalize_upload_list(investor_uploads)
    concall_files = normalize_upload_list(concall_uploads)
    # Store screenshots in session state for post-extraction processing
    st.session_state["annual_screenshots_data"] = annual_screenshots or []
    st.session_state["investor_screenshots_data"] = investor_screenshots or []

    ann_imgs = len(annual_screenshots) if annual_screenshots else 0
    inv_imgs = len(investor_screenshots) if investor_screenshots else 0
    if annual_files or investor_files or concall_files or ann_imgs or inv_imgs:
        st.caption(
            f"Ready: {len(annual_files)} annual PDF(s), "
            f"{len(investor_files)} investor PDF(s), "
            f"{len(concall_files)} concall PDF(s)"
            + (f", {ann_imgs} annual screenshot(s)" if ann_imgs else "")
            + (f", {inv_imgs} investor screenshot(s)" if inv_imgs else "")
            + "."
        )

    st.divider()
    st.subheader("AI Extraction (Recommended)")
    st.caption(
        "Gemini reads your PDFs visually for financial metrics — banks, NBFCs, HFCs. "
        "pdfplumber fills any gaps. Free key from aistudio.google.com."
    )
    default_gemini = ""
    try:
        default_gemini = st.secrets.get("GEMINI_API_KEY", "")
    except Exception:
        pass
    gemini_key_upload = st.text_input(
        "Gemini API Key",
        value=default_gemini,
        type="password",
        placeholder="AQ.... or AIza...",
        help="Primary PDF extraction via Gemini File API. pdfplumber runs afterward for missing cells.",
        key="gemini_api_key_upload",
    )
    st.session_state["gemini_api_key"] = gemini_key_upload
    if gemini_key_upload == default_gemini and default_gemini:
        st.caption(
            "Using default Gemini key from secrets. "
            "Replace with your own from aistudio.google.com if needed."
        )
    elif not gemini_key_upload:
        st.caption("Optional — without a key, pdfplumber-only extraction still runs.")

    run_clicked = st.button("Run Extraction", type="primary")

    if not run_clicked:
        st.info("Upload PDFs and click **Run Extraction** to scan and extract metrics.")
        return

    missing = []
    has_annual = bool(annual_files or ann_imgs)
    has_investor = bool(investor_files or inv_imgs)
    if not has_annual and not has_investor:
        missing.append(
            "at least one Annual Report or Investor Presentation (PDF or screenshot)"
        )
    if missing:
        st.error("Please upload: " + ", ".join(missing))
        return

    ok = run_full_pipeline(annual_files, investor_files, concall_files)
    if ok:
        st.rerun()


# -----------------------------------------------------------------------------
# Top-level controller
# -----------------------------------------------------------------------------
def render_sidebar() -> None:
    with st.sidebar:
        st.header("Session")
        if SS_REVIEWED in st.session_state:
            st.markdown("Status: **Extraction loaded**")
            if st.session_state.get(SS_APPROVED):
                st.markdown("Approval: **Approved**")
            else:
                st.markdown("Approval: **Pending review**")
            if st.button("Clear & re-upload"):
                reset_pipeline()
                st.rerun()
        else:
            st.markdown("Status: **No extraction yet**")
        if st.session_state.get(SS_PHASE6_DONE):
            st.markdown("Enterprise report: **Generated**")
        elif st.session_state.get(SS_REPORT_DONE):
            st.markdown("Analytical report: **Generated**")
        elif st.session_state.get(SS_COMMENTARY_DONE):
            st.markdown("Commentary: **Done**")
        st.caption("Phases 1–6 active.")


def main() -> None:
    st.set_page_config(
        page_title="Credit Review Report Generator",
        page_icon="📊",
        layout="wide",
    )

    render_sidebar()

    if SS_REVIEWED in st.session_state and st.session_state[SS_REVIEWED]:
        render_review_workflow()
    else:
        render_upload_screen()


if __name__ == "__main__":
    main()
